import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from typing import Dict, Any, List

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set margins/padding of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_process_document(diagram_title: str, description: str, react_flow_data: Dict[str, Any], notes: Dict[str, str], powl_code: str = "", simplified: bool = False) -> io.BytesIO:
    """
    Generates a premium styled Word (.docx) document summarizing the mapped process.
    If simplified=True, it generates a shorter version focusing on Flow, Systems, and Areas.
    """
    doc = Document()
    
    # 1. Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Define Colors
    c_primary = RGBColor(0x0f, 0x17, 0x2a)   # Slate 900
    c_secondary = RGBColor(0x47, 0x55, 0x69) # Slate 600
    c_accent = RGBColor(0x02, 0x84, 0xc7)    # Sky 600
    c_text = RGBColor(0x33, 0x41, 0x55)      # Slate 700

    # Configure Default Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = c_text

    # Title Page/Heading
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    
    title_run = title_p.add_run(f"Documentação do Processo: {diagram_title}")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = c_primary

    # Subtitle / Metadata
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(24)
    meta_run = meta_p.add_run("Gerado automaticamente pelo Cresol Process Mining Platform")
    meta_run.font.italic = True
    meta_run.font.color.rgb = c_secondary
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: Descrição Geral
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1_run = h1.add_run("1. Descrição Geral do Processo")
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = c_accent

    desc_p = doc.add_paragraph(description or "Nenhuma descrição fornecida.")
    desc_p.paragraph_format.line_spacing = 1.15
    desc_p.paragraph_format.space_after = Pt(16)

    # Section 2: Raias e Responsabilidades (Lanes/Pools)
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2_run = h2.add_run("2. Estrutura de Responsabilidades (Swimlanes)")
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = c_accent

    # Extract nodes to list pools and lanes
    nodes = react_flow_data.get('nodes', [])
    swimlanes_layout = react_flow_data.get('swimlanes', [])
    
    if swimlanes_layout:
        doc.add_paragraph("O processo está mapeado nas seguintes áreas e raias de atuação:")
        
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Format Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Pool (Entidade/Organização)"
        hdr_cells[1].text = "Lane (Papel/Departamento)"
        
        for cell in hdr_cells:
            set_cell_background(cell, "0f172a") # Dark Slate
            set_cell_margins(cell, 120, 120, 150, 150)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

        for lane_info in swimlanes_layout:
            row_cells = table.add_row().cells
            row_cells[0].text = lane_info.get('pool', '')
            row_cells[1].text = lane_info.get('lane', '')
            for cell in row_cells:
                set_cell_margins(cell, 100, 100, 150, 150)
                set_cell_background(cell, "f8fafc") # light background
    else:
        doc.add_paragraph("Não foram identificadas subdivisões de raias (swimlanes) específicas neste processo.")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Detalhamento das Atividades e Anotações
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)
    h3_run = h3.add_run("3. Detalhamento das Atividades")
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = c_accent

    task_nodes = [n for n in nodes if n.get('type') == 'task']
    
    if task_nodes:
        if simplified:
            # Documento Simplificado: Focar apenas no Fluxo, Sistemas e Variáveis
            h3_run.text = "3. Relação de Sistemas e Documentos do Fluxo"
            
            task_table = doc.add_table(rows=1, cols=3)
            task_table.autofit = False
            widths = [Inches(2.5), Inches(2.0), Inches(2.0)]
            
            hdr_cells = task_table.rows[0].cells
            hdr_cells[0].text = "Atividade"
            hdr_cells[1].text = "Sistemas Envolvidos"
            hdr_cells[2].text = "Variáveis / Docs"
            
            for i, cell in enumerate(hdr_cells):
                cell.width = widths[i]
                set_cell_background(cell, "0284c7")
                set_cell_margins(cell, 120, 120, 150, 150)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            
            for task in task_nodes:
                label = task['data'].get('label', '')
                sys_list = task['data'].get('systems', [])
                if isinstance(sys_list, str):
                    sys_list = [s.strip() for s in sys_list.split(',')]
                var_list = task['data'].get('variables', [])
                if isinstance(var_list, str):
                    var_list = [s.strip() for s in var_list.split(',')]
                
                row_cells = task_table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = ", ".join(sys_list) if sys_list else "-"
                row_cells[2].text = ", ".join(var_list) if var_list else "-"
                
                for i, cell in enumerate(row_cells):
                    cell.width = widths[i]
                    set_cell_margins(cell, 100, 100, 150, 150)
                    set_cell_background(cell, "f8fafc")
        else:
            # Documento Completo
            # Create table of tasks
            task_table = doc.add_table(rows=1, cols=4)
            task_table.autofit = False
        
        # Set widths roughly
        widths = [Inches(1.8), Inches(1.2), Inches(1.2), Inches(2.3)]
        
        hdr_cells = task_table.rows[0].cells
        hdr_cells[0].text = "Atividade"
        hdr_cells[1].text = "Pool"
        hdr_cells[2].text = "Lane"
        hdr_cells[3].text = "Observações / Regras de Negócio"
        
        for i, cell in enumerate(hdr_cells):
            cell.width = widths[i]
            set_cell_background(cell, "0284c7") # Sky 600
            set_cell_margins(cell, 120, 120, 150, 150)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

        for task in task_nodes:
            label = task['data'].get('label', '')
            pool = task['data'].get('pool', '')
            lane = task['data'].get('lane', '')
            
            # Fetch note either from frontend node state or DB notes
            note_content = task['data'].get('annotations', '')
            if not note_content and notes:
                # Try fallback matching in DB notes dict
                note_content = notes.get(label, '') or notes.get(task['id'], '')

            row_cells = task_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = pool
            row_cells[2].text = lane
            row_cells[3].text = note_content or "-"
            
            for i, cell in enumerate(row_cells):
                cell.width = widths[i]
                set_cell_margins(cell, 100, 100, 150, 150)
                # alternating row backgrounds for readability
                set_cell_background(cell, "f8fafc")
    else:
        doc.add_paragraph("Nenhuma atividade de processo estruturada foi encontrada.")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 4: Estrutura Lógica do Processo (POWL)
    if not simplified and powl_code:
        h4 = doc.add_paragraph()
        h4.paragraph_format.space_before = Pt(18)
        h4.paragraph_format.space_after = Pt(8)
        h4_run = h4.add_run("4. Estrutura Lógica (POWL Code)")
        h4_run.font.size = Pt(16)
        h4_run.font.bold = True
        h4_run.font.color.rgb = c_accent

        doc.add_paragraph("O modelo matemático do processo (representação em código estruturado) gerado pelos agentes:")
        
        # Add code block with gray background
        code_p = doc.add_paragraph()
        code_p.paragraph_format.left_indent = Inches(0.2)
        code_p.paragraph_format.right_indent = Inches(0.2)
        code_p.paragraph_format.space_before = Pt(6)
        code_p.paragraph_format.space_after = Pt(6)
        
        # We can simulate a shaded background code snippet in word by setting paragraph borders/shading
        pPr = code_p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f1f5f9"/>')
        pPr.append(shd)
        
        # Border
        pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="cbd5e1"/></w:pBdr>')
        pPr.append(pbdr)

        code_run = code_p.add_run(powl_code)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9.5)
        code_run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
